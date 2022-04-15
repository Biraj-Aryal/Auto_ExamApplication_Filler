from docxtpl import DocxTemplate
from datetime import date
import docx
from docx.shared import Pt
from pyBSDate import convert_AD_to_BS
from tkinter import *
from tkinter.ttk import *
from tkinter.filedialog import askopenfile
import os
import sqlite3

directory = os.getcwd()
# logo_path = os.path.join(os.path.join(directory, 'app_resources'), 'quizzy.ico')
# template_path = os.path.join(os.path.join(directory, 'app_resources'), "template.docx")
# database_path = os.path.join(os.path.join(directory, 'app_resources'), 'up_date.db')

logo_path = 'quizzy.ico'
template_path = "template.docx"
database_path = 'up_date.db'

conn = sqlite3.connect(database_path)
c = conn.cursor()

# sq lite part
# def create_table_examinees():
#     with conn:
#         c.execute("""CREATE TABLE examinees (
#         reg integer,
#         year integer,
#         sem integer,
#         ef_n text,
#         em_n text,
#         el_n text,
#         nf_n text,
#         n_mn text,
#         n_ln text,
#         AD_year integer,
#         AD_month integer,
#         AD_day integer,
#         gender text
#         )""")
#
#
# def delete_table(table):
#     with conn:
#         c.execute(f"DROP TABLE {table}")

def is_update_created():
    with conn:
        c.execute("SELECT * FROM examinees")

        if len(c.fetchall()) < 1:
            return False
        else:
            return True

# gets input from user and stores it in the database; replaces old data.
def input_and_update():
    def first_update():
        with conn:
            c.execute("INSERT INTO examinees VALUES ("
                      "?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                      (regis, usr_year, usr_sem, first_n, middle_n, last_n, nepali_fname, nepali_mname, nepali_lname, year_e, month_e, day_e, usr_gender))


    def further_update():
        with conn:
            c.execute("""
            UPDATE examinees
            SET reg = ?,
            year = ?,
            sem = ?,
            ef_n = ?,
            em_n = ?,
            el_n = ?,
            nf_n = ?,
            n_mn = ?,
            n_ln = ?,
            AD_year = ?,
            AD_month = ?,
            AD_day = ?,
            gender = ?;""", (regis, usr_year, usr_sem, first_n, middle_n, last_n, nepali_fname, nepali_mname, nepali_lname, year_e, month_e, day_e, usr_gender))

    regis = kinterreg
    usr_year = kinteryear
    usr_sem = kintersem
    first_n = kinterfne
    middle_n = kintermne
    last_n = kinterlne
    nepali_fname = kinterfnn
    nepali_mname = kintermnn
    nepali_lname = kinterlnn
    year_e = kinterdoby
    month_e = kinterdobm
    day_e = kinterdobd
    usr_gender = kintergender

    if is_update_created():
        further_update()

    elif not is_update_created():
        first_update()

# just testing what's in the table before running the operation
def get_table():
    with conn:
        c.execute(f"SELECT * FROM examinees")
        return c.fetchone()

root = Tk()
# root.geometry('300x200')
root.title('Quizzy Fill')
root.iconbitmap(logo_path)
root.configure(bg='khaki')


def fill_my_form():
    # All about time
    today = date.today()
    today = today.strftime("%B %d, %Y")
    this_year = today.split(',')[1][1:]
    year_one = str(int(this_year) - 1)[-2:]
    year_two = this_year[-2:]


    #Loading the template via docxtemplate modules
    doc = DocxTemplate(template_path)


    # !!!!!!!   year section
    usr_year = int(kinteryear)
    if usr_year == 1:
        usr_year_s = 'st'
    elif usr_year == 2:
        usr_year_s = 'nd'
    elif usr_year == 3:
        usr_year_s = 'rd'
    else:
        usr_year_s = 'th'

    # !!!!!!!!    sem section
    usr_sem = int(kintersem)
    if usr_sem == 1:
        usr_sem_s = 'st'
    elif usr_sem == 2:
        usr_sem_s = 'nd'

    # !!!!!!!!!  reg section
    regis = kinterreg
    batcher = '20' + regis[-2:]

    # !!!!!!!!!!   name section
    first_n = kinterfne.strip()
    middle_n = kintermne.strip()
    last_n = kinterlne.strip()

    # !!!!!!!! location of the photo and signature (Create a browser to select the images)
    ph_locate = photo_location
    sgn_locate = sign_location

    doc.replace_pic("dummy_photo.png", ph_locate)
    doc.replace_pic("dummy_signature.png", sgn_locate)

    # subjects section: It will depend on what the usr_year and usr_sem
    if usr_year == 1 and usr_sem == 1:
        subject = {1: ['English I ( General)', 'LENG', '101'],
                   2: ['Jurisprudence', 'LJUR', '111'],
                   3: ['Constitutional Law I', 'LCNS', '121'],
                   4: ['Contract Law', 'LCON', '161'],
                   5: ['Financial Accounting', 'ACC', '205'],
                   6: ['Business Management', 'GEM', '231']}
    if usr_year == 1 and usr_sem == 2:
        subject = {1: ['Microeconomics', 'ECO', '201'],
                   2: ['Nepali (General)', 'LNEP', '103'],
                   3: ['Competition Law', 'LCMP', '162'],
                   4: ['English II (Legal English)', 'LENG', '102'],
                   5: ['Quantitative Techniques', 'MAS', '103'],
                   6: ['Constitutional Law II', 'LCNS', '122']}
    if usr_year == 2 and usr_sem == 1:
        subject = {1: ['Managerial Communication', 'GEM', '201'],
                   2: ['Nepali II (Legal Language)', 'LNEP', '204'],
                   3: ['Management Accounting', 'ACC', '210'],
                   4: ['Human Rights Law and Practice', 'LHRT', '223'],
                   5: ['Torts & Consumer Protection Laws', 'LCPT', '231'],
                   6: ['Family Law', 'LFAM', '232']}
    if usr_year == 2 and usr_sem == 2:
        subject = {1: ['Macroeconomics', 'ECO', '210'],
                   2: ['Legal Reasoning Skill & Logic', 'LRSK', '233'],
                   3: ['Law of Crimes I (Penal Code)', 'LCRM', '241'],
                   4: ['Organizational Behavior', 'HRM', '320'],
                   5: ['Cyber Law', 'LCYB', '234'],
                   6: ['Philosophy of Life & Life Style', 'LPHI', '205']}
    if usr_year == 3 and usr_sem == 1:
        subject = {1: ['Human Resource Management', 'HRM', '201'],
                   2: ['Financial Management', 'FIN', '301'],
                   3: ['Property Law', 'LPRT', '335'],
                   4: ['Law of Crimes II (Criminal Procedure Code I)', 'LCRM', '341'],
                   5: ['Company Law I', 'LCOM', '351'],
                   6: ['Administration Law', 'LADM', '336']}
    if usr_year == 3 and usr_sem == 2:
        subject = {1: ['Environmental law', 'LENV', '371'],
                   2: ['Intellectual Property Law', 'LINP', '368'],
                   3: ['Marketing Management', 'MKT', '310'],
                   4: ['Company Law II', 'LCOM', '352'],
                   5: ['Law of Crimes II', 'LCRM', '342'],
                   6: ['International Business', 'GEM', '470']}
    if usr_year == 4 and usr_sem == 1:
        subject = {1: ['Project Management', 'GEM', '332'],
                   2: ['Banking Law', 'LBNK', '453'],
                   3: ['Public International Law', 'LINT', '472'],
                   4: ['Civil Procedure & Limitation Law I', 'LPCD', '437'],
                   5: ['Financial Institutions & Markets', 'FIN', '330'],
                   6: ['Merger & Acquisitions Law', 'LMRG', '454']}
    if usr_year == 4 and usr_sem == 2:
        subject = {1: ['Civil Procedure and Limitation Law II', 'LPCD', '439'],
                   2: ['Taxation Law', 'LTAX', '464'],
                   3: ['Corporate Governance and Business Ethics', 'LCGB', '481'],
                   4: ['Law of Evidence', 'LEVD', '438'],
                   5: ['Strategic Management', 'GEM', '490'],
                   6: ['Insurance Law', 'LINS', '463']}
    if usr_year == 5 and usr_sem == 1:
        subject = {1: ['Entrepreneurship and New Business Formation', 'GEM', '310'],
                   2: ['Water and Energy Law', 'LWTR', '573'],
                   3: ['Labour and Industrial Law', 'LIND', '565'],
                   4: ['Investment Law', 'LINV', '566'],
                   5: ['Legal Research Analysis and Writing', 'LDBS', '506'],
                   6: ['Trade Law', 'LTRD', '567']}
    if usr_year == 5 and usr_sem == 2:
        subject = {1: ['Drafting, Pleading and Conveyance', 'LDPC', '591'],
                   2: ['Alternative Dispute Resolution', 'LADR', '592'],
                   3: ['Professional Ethics and Professional Accounting System', 'LPEA', '593'],
                   4: ['Moot Court Exercise and Internship', 'LMTC', '594'],
                   5: ['Court Planning and Management', 'LCPM', '595'],
                   6: ['\n', '', '']}

    # fix for the 10th sem
    if usr_year == 5 and usr_sem == 2:
        fix = ''
    else:
        fix = '3'

    context = {
        # the subjects
        'one': f'{subject[1][0]}',
        'two': f'{subject[2][0]}',
        'three': f'{subject[3][0]}',
        'four': f'{subject[4][0]}',
        'five': f'{subject[5][0]}',
        'six': f'{subject[6][0]}',
        'one_c1': f'{subject[1][1]}',
        'two_c1': f'{subject[2][1]}',
        'three_c1': f'{subject[3][1]}',
        'four_c1': f'{subject[4][1]}',
        'five_c1': f'{subject[5][1]}',
        'six_c1': f'{subject[6][1]}',
        'one_c2': f'{subject[1][2]}',
        'two_c2': f'{subject[2][2]}',
        'three_c2': f'{subject[3][2]}',
        'four_c2': f'{subject[4][2]}',
        'five_c2': f'{subject[5][2]}',
        'six_c2': f'{subject[6][2]}',

        # reg, name, batch, year, sem, today's date
        'reg' : f'{regis}',
        'f_name': f'{first_n}',
        'm_name': f'{middle_n}',
        'l_name': f'{last_n}',
        'batch': f'{batcher}',
        'year_s' : f'{usr_year}',
        'year_super': f'{usr_year_s}',
        'sem_s': f'{usr_sem}',
        'sem_super': f'{usr_sem_s}',
        'today_date': f'{today}',
        'y_1': f'{year_one}',
        'y_2': f'{year_two}',
        'fix_10': f'{fix}'}

    # Saving the channges
    doc.render(context)
    doc.save(r"exam_form.docx")

    # Using the next module
    d = docx.Document(r"exam_form.docx")

    the_name = last_n.upper() + ' ' + first_n.upper() + ' ' + middle_n.upper()
    the_name = the_name.strip()

    def fix_font(run):
        font = run.font
        select_font = 'Times New Roman'
        font.name = select_font
        font.size = Pt(12)

    # ????? here, if len(the_name) is above certain number, maybe throw an error
    for index, ch in enumerate(the_name):
        d.tables[6].cell(0, index).text = ch
        try:
            c = d.tables[6].cell(0, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass

    # !!!!!!!   All about the Name in Nepali (pg 2)
    nepali_fname = kinterfnn
    nepali_mname = kintermnn
    nepali_lname = kinterlnn
    nepali_fname = nepali_fname.split(' ')
    nepali_mname = nepali_mname.split(' ')
    nepali_lname = nepali_lname.split(' ')

    f = the_name.split(first_n.upper())

    # determination of space in name having middle name
    if len(f[1]) > 1:
        # Space after last name
        first_space = 0
        second_space = 0
        for i in f[0]:
            if i == ' ':
                first_space += 1
        # Space after first name
        for i in f[1]:
            if i == ' ':
                second_space += 1

    # determination of space in name not having middle name
    elif len(f[1]) < 2:
        # Space after last name
        first_space = 0
        second_space = 0
        for i in f[0]:
            if i == ' ':
                first_space += 1

    # determination of space in name which is somehow messed up
    else:
        if nepali_mname == '':
            first_space = 6
            second_space = 3
        else:
            first_space = 5
            second_space = 3

    first_space = len(last_n) + first_space - len(nepali_lname)
    second_space = len(first_n) + second_space - len(nepali_fname)

    def int_to_list(number):
        my_list = []
        for i in range(number):
            my_list.append('')
        return my_list


    first_space = int_to_list(first_space)
    second_space = int_to_list(second_space)

    nepali = nepali_lname + first_space + nepali_fname + second_space + nepali_mname

    for index, ch in enumerate(nepali):
        d.tables[6].cell(1, index).text = ch
        try:
            c = d.tables[6].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass

    year_e = kinterdoby
    month_e = kinterdobm
    day_e = kinterdobd

    # The input parts are the three variables above.

    # messy stuff part
    to_convert_year = int(year_e)
    to_convert_month = int(month_e)
    to_convert_day = int(day_e)

    if len(month_e) == 1:
        month_e = '0' + month_e

    if len(day_e) == 1:
        day_e = '0' + day_e

    DOB_E = year_e + month_e + day_e
    for index, ch in enumerate(DOB_E):
        index += 9
        d.tables[7].cell(1, index).text = ch
        try:
            c = d.tables[7].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass

    bs_date = convert_AD_to_BS(to_convert_year, to_convert_month, to_convert_day)
    year_n, month_n, day_n = str(bs_date[0]), str(bs_date[1]), str(bs_date[2])

    if len(month_n) == 1:
        month_n = '0' + month_n

    if len(day_n) == 1:
        day_n = '0' + day_n

    DOB_N = year_n + month_n + day_n
    for index, ch in enumerate(DOB_N):
        d.tables[7].cell(1, index).text = ch
        try:
            c = d.tables[7].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass


    # Block letter registration
    for index, ch in enumerate(regis):
        d.tables[4].cell(0, index).text = ch
        try:
            c = d.tables[7].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass


    # !!! Give a scroll of the options of 'tick on Male', 'tick on Female'. Maybe add an option of 'don't tick'.
    usr_gender = kintergender.lower()

    p = d.paragraphs[30]
    if usr_gender in ['m', 'male', 'tick male']:
        p.add_run(':  Male✔       Female')

    elif usr_gender in ['f', 'female', 'tick female']:
        p.add_run(':  Male        Female✔')

    else:
        p.add_run(':  Male        Female')

    c = d.paragraphs[30].runs[1]
    fix_font(c)

    d.save(f"exam_form.docx")

def tkinter_interface(y, m, d, gender, reg_n, year, sem, rf, rm, rl, df, dm, dl):
    # DOB (AD): Label
    DOBlabel = Label(root, text="Date of Birth (in AD):", font=('Times_New_Roman', 20))
    DOBlabel.grid(row=2, column=0, columnspan=2, padx=(95, 0), pady=(30,0))


    # DOB: Choices years
    year_choices = ['1990', '1991', '1992', '1993', '1994', '1995', '1996', '1997', '1998', '1999', '2000', '2001', '2002',
                     '2003', '2004', '2005', '2006', '2007', '2008', '2009', '2010']
    yearbox = Combobox(root, values=year_choices, width=4, font=('Times_New_Roman', 16))
    yearbox.set(y)
    yearbox['state'] = 'readonly'  # normal
    yearbox.grid(row=2, column=2, padx=(0,20), pady=(30,0))

    # DOB: Choices month
    month_choices = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']
    monthbox = Combobox(root, values=month_choices, width=2, font=('Times_New_Roman', 16))
    monthbox.set(m)
    monthbox['state'] = 'readonly'  # normal
    monthbox.grid(row=2, column=3, padx=(0, 23), pady=(30,0))

    # DOB: Choices days
    day_choices = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                   '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
    daybox = Combobox(root, values=day_choices, width=2, font=('Times_New_Roman', 16))
    daybox.set(d)
    daybox['state'] = 'readonly'  # normal
    daybox.grid(row=2, column=4, padx=(0, 170), pady=(30,0))

    # Gender: Label
    Genderlabel = Label(root, text="Gender:", font=('Times_New_Roman', 20))
    Genderlabel.grid(row=3, column=1, padx=(88, 0))

    # Gender: Choices entry
    gender_choices = ['Tick Male', 'Tick Female', "Empty"]
    # selected_gender = StringVar(root)
    genderbox = Combobox(root, values=gender_choices, width=9, font=('Times_New_Roman', 16))
    genderbox.set(gender)
    genderbox['state'] = 'readonly'  # normal
    genderbox.grid(row=3, column=2, columnspan=2, padx=(0, 49))

    # Reg NO
    Reglabel = Label(root, text="Registration No:", font=('Times_New_Roman', 20))
    Reglabel.grid(row=2, column=6, columnspan=1, padx=(50, 0), pady=(30, 0))

    rege = Entry(root, width=12, font=('Times_New_Roman', 16))
    rege.grid(row=2, column=7, columnspan=3, padx=(0, 110), pady=(30, 0))
    rege.insert(0, reg_n)

    # Year and Sem
    # Year
    YearLabel = Label(root, text="Year:", font=('Times_New_Roman', 20))
    YearLabel.grid(row=3, column=6, padx=(115, 0))

    years_choices = ['1', '2', '3', '4', '5']
    YearsBox = Combobox(root, values=years_choices, width=1, font=('Times_New_Roman', 16))
    YearsBox.set(year)
    YearsBox['state'] = 'readonly'  # normal
    YearsBox.grid(row=3, column=7, padx=(0, 0))

    # Sem
    SemLabel = Label(root, text="Sem:", font=('Times_New_Roman', 20))
    SemLabel.grid(row=3, column=8)

    Sem_choices = ['1', '2']
    SemBox = Combobox(root, values=Sem_choices, width=1, font=('Times_New_Roman', 16))
    SemBox.set(sem)
    SemBox['state'] = 'readonly'  # normal
    SemBox.grid(row=3, column=9, padx=(0, 120))

    # Roman Name Note
    RMLabel = Label(root, text="Replace the following name with your name in English. Only capaitalize (the initials) as shown.")
    RMLabel.grid(row=4, column=0, columnspan=8, pady=(65, 5), padx=(105, 0))

    # Roman first name
    RFlabel = Label(root, text="First Name:", font=('Times_New_Roman', 20))
    RFlabel.grid(row=5, column=0, padx=(0, 0))

    fnamee = Entry(root, width=15, font=('Times_New_Roman', 16))
    fnamee.grid(row=5, column=1, padx=(0, 0))
    fnamee.insert(0, rf)

    #Roman middle name
    RMlabel = Label(root, text="Middle Name:", font=('Times_New_Roman', 20))
    RMlabel.grid(row=5, column=2, columnspan=2, padx=(40, 0))

    mnamee = Entry(root, width=15, font=('Times_New_Roman', 16))
    mnamee.grid(row=5, column=4, columnspan=2, padx=(0, 40))
    mnamee.insert(0, rm)

    #Roman last name
    RLlabel = Label(root, text="Last Name:", font=('Times_New_Roman', 20))
    RLlabel.grid(row=5, column=6, padx=(80, 0))
    #
    lnamee = Entry(root, width=15, font=('Times_New_Roman', 16))
    lnamee.grid(row=5, column=7, columnspan=3, padx=(0, 80))
    lnamee.insert(0, rl)


    # Devanagari Name Note
    DMLabel = Label(root, text="Replace the following with your name in Nepali. You ought to put a space Between each characters as shown.")
    DMLabel.grid(row=6, column=0, columnspan=9, pady=(70, 5), padx=(65, 0))

    # Devanagari first name
    DFlabel = Label(root, text="First Name:", font=('Times_New_Roman', 20))
    DFlabel.grid(row=7, column=0, padx=(0, 0))

    fnamen = Entry(root, width=15, font=('Times_New_Roman', 16))
    fnamen.grid(row=7, column=1, padx=(0, 0))
    fnamen.insert(0, df)

    #Devanagari middle name
    NMlabel = Label(root, text="Middle Name:", font=('Times_New_Roman', 20))
    NMlabel.grid(row=7, column=2, columnspan=2, padx=(45, 0))

    mnamen = Entry(root, width=15, font=('Times_New_Roman', 16))
    mnamen.grid(row=7, column=4, columnspan=2, padx=(0, 45))
    mnamen.insert(0, dm)

    #Devanagari last name
    NLlabel = Label(root, text="Last Name:", font=('Times_New_Roman', 20))
    NLlabel.grid(row=7, column=6, padx=(80, 0))
    #
    lnamen = Entry(root, width=15, font=('Times_New_Roman', 16))
    lnamen.grid(row=7, column=7, columnspan=3, padx=(0, 80))
    lnamen.insert(0, dl)

    # # Middle Name Note
    # DMLabel = Label(root, text="(If middle name is not applicable, make sure to empty the fields.)")
    # DMLabel.grid(row=8, column=0, columnspan=9, pady=(30,5), padx=(0,410))

    # # Browser Note
    # DMLabel = Label(root, text="Browse to select images of your photo and your signature below.", font="Raleway")
    # DMLabel.grid(row=9, column=0, columnspan=9, pady=(50,5))

    file_types = [("PNG file", "*.png"), ("EMG file", "*.emg"), ("JPG file", "*.jpg"), ("JPG file2", "*.jpeg"),
    ("WMF file", "*.wmf"), ("TIFF file", "*.tif"), ("TIFF file2", "*.tiff")]

    def photo_browser():
        file = askopenfile(parent=root, mode='rb', title="Choose a file", filetypes=file_types)
        if file:
            global photo_location
            photo_location = str(file).split('name=')[-1][:-1].strip("'")
            file_name = os.path.basename(photo_location)
            photoe.configure(state='normal')
            photoe.delete(0, "end")
            photoe.insert(0, file_name.rstrip("'"))
            photoe.configure(state='readonly')

    def signature_browser():
        file = askopenfile(parent=root, mode='rb', title="Choose a file", filetypes=file_types)
        if file:
            global sign_location
            sign_location = str(file).split('name=')[-1][:-1].strip("'")
            file_name = os.path.basename(sign_location)
            signe.configure(state='normal')
            signe.delete(0, "end")
            signe.insert(0, file_name.rstrip("'"))
            signe.configure(state='readonly')

    def generate_form():
        # Names in both languages
        global kinterfne, kintermne, kinterlne, kinterfnn, kintermnn, kinterlnn, kinterreg, kinterdoby, kinterdobm, kinterdobd, kinteryear, kintersem, kintergender
        kinterfne = fnamee.get()
        kintermne = mnamee.get()
        kinterlne = lnamee.get()
        kinterfnn = fnamen.get()
        kintermnn = mnamen.get()
        kinterlnn = lnamen.get()
        #Registration Number
        kinterreg = rege.get()
        # DOB
        kinterdoby = yearbox.get()
        kinterdobm = monthbox.get()
        kinterdobd = daybox.get()
        # Year and Sem
        kinteryear = YearsBox.get()
        kintersem = SemBox.get()
        # Gender
        kintergender = genderbox.get()
        input_and_update()  # Sqlite3 update
        fill_my_form() # docx stuff
        # Ending note regarding location of the file
        Notif_Label = Label(root, text=f"The form is in the same folder as this program: {os.path.abspath(directory)}", font=('Times_New_Roman', 15))
        Notif_Label.grid(row=11, column=0, columnspan=9, pady=(20, 10))
        conn.close()
        root.after(6200, lambda: root.destroy())


    # Browse your photo
    browse_photo_btn = Button(root, text="\nSelect your photo\n", command=photo_browser)
    browse_photo_btn.grid(column=1, row=10, columnspan=3, padx=(0,90), pady=(50, 0))

    photoe = Entry(root, width=20)
    photoe.grid(row=11, column=1, columnspan=3, pady=(0,120), padx=(0,90))
    photoe.insert(0, "No image has been selected")
    photoe.configure(state='readonly')

    # Browse your signature
    browse_signature_btn = Button(root, text="\nSelect your signature\n", command=signature_browser)
    browse_signature_btn.grid(column=5, row=10, columnspan=4, padx=(0,0), pady=(50, 0))

    signe = Entry(root, width=20)
    signe.grid(row=11, column=5, columnspan=4, pady=(0,120), padx=(10,0))
    signe.insert(0, "No image has been selected")
    signe.configure(state='readonly')

    # #Generate Button
    generate_btn = Button(root, text="\nGenerate Form\n", command=generate_form)
    generate_btn.grid(column=3, row=12, columnspan=2, pady=(0, 30), padx=(15,0))

my_table = get_table()
print(my_table)

if not is_update_created():
    print('Update is not created')
    tkinter_interface('2000', '9', '30', 'Empty', "123456-78", '3', '2', "Abdul", "Magomedov", "Nurmagomed", "दा स", "ह रि", "मा ध व")

else:
    print('Update is created')
    tkinter_interface(str(my_table[9]), str(my_table[10]), str(my_table[11]), str(my_table[12]), str(my_table[0]), str(my_table[1]), str(my_table[2]), str(my_table[3]), str(my_table[4]), str(my_table[5]), str(my_table[6]),
                      str(my_table[7]), str(my_table[8]))

root.mainloop()
